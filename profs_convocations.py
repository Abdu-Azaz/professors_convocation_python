import tkinter as tk
from tkinter import filedialog, messagebox
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
from docx.shared import Pt

today = datetime.date.today()

def read_excel_file():
    """
    Open file dialog to select Excel file.
    """
    file_path = filedialog.askopenfilename()
    return file_path

def process_excel(file_path, session, period, au):
    """
    Process Excel file and generate DOCX document.
    """
    raw_data_df = pd.read_excel(file_path, skiprows=3, header=None, index_col=[0], engine='openpyxl')
    processed_data_df = pretreatment(raw_data_df)
    schedule_data = grouping_profs_info_in_a_dict(raw_data_df, processed_data_df)
    generate_docx(schedule_data, session, period, au)

def pretreatment(raw_data_df):
    """
    Perform pretreatment on the DataFrame.
    """
    raw_data_df.iloc[7:, 0] = raw_data_df.iloc[7:, 0].ffill()
    raw_data_df.iloc[0, :] = raw_data_df.iloc[0, :].ffill()
    raw_data_df.iloc[1, :] = raw_data_df.iloc[1, :].ffill()
    # raw_data_df.to_excel("raw_data_df.xlsx", index=False)
    return raw_data_df.iloc[6:,:]

def grouping_profs_info_in_a_dict(raw_data_df, processed_data_df):
    """
    Group information about professors' schedules in a dictionary.
    """
    schedule_data = {}
    for x, row in enumerate(processed_data_df.values):  
        for y, value in enumerate(row[1:-2], start=1):    
            if not pd.isna(value) and  value[-1].isdigit():  # Check if the cell is not NaN
                professor = processed_data_df.iloc[x, 0]
                date = raw_data_df.iloc[0, y]
                time = raw_data_df.iloc[1, y]
                niveau = raw_data_df.iloc[3, y]
                subject = raw_data_df.iloc[4, y]
                local = value  # Cell value representing the exam location
                schedule_info = {'subject': subject, 'date': date, 'time': time, 'niveau': niveau, 'local': local}
                if professor in schedule_data:
                    schedule_data[professor].append(schedule_info)
                else:
                    schedule_data[professor] = [schedule_info]
    return schedule_data


def generate_docx(schedule_data, session, period, au):
    """
    Generate DOCX document from schedule data.
    """
    doc = Document()
    default_font = doc.styles['Normal'].font
    default_font.size = Pt(12)
    default_font.name = 'Arial'
    

    for professor, info_list in schedule_data.items():
        logo_paragraph = doc.add_paragraph()
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo_path, width=Inches(3.3))
        # logo_run.add_picture('stand_fr_dark.png', width=Inches(3.3))
        logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()
        doc.add_paragraph(f"A Mme/ Mr: {professor}")
        doc.add_paragraph(f"Objet: Convocation aux surveillances des Examens")
        doc.add_paragraph("Cher(e) collègue,")
        doc.add_paragraph(f"Nous vous saurions gé de bien vouloir prendre toutes les dispositions nécessaires pour assurer la surveillance des épreuves écrites de la session {session} de {period} {au} aux jours et horaires indiqués ci-dessus:")
        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Module'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Horaire'
        hdr_cells[3].text = 'Niveau'
        hdr_cells[4].text = 'Local'
        #
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    
        #
        table.allow_autofit = False  
        for info in info_list:
            subject = str(info['subject'])
            date = str(info['date'])
            time = str(info['time'])
            niveau = str(info['niveau'])
            local = str(info['local'])
            row_cells = table.add_row().cells
            row_cells[0].text = subject
            row_cells[1].text = date
            row_cells[2].text = time
            row_cells[3].text = niveau
            row_cells[4].text = local
        doc.add_paragraph()
        
        doc.add_paragraph("Nous vous remercions de votre précieuse collaboration")
        doc.add_paragraph()
        doc.add_paragraph()
        
        doc.add_paragraph(f"Ait Melloul le: {today.strftime('%d-%m-%Y')} ").alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()
        doc.add_paragraph(f"Le doyen").alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_page_break()

    filename = f"invitations_profs_{today.strftime('%Y_%m_%d')}.docx"
    doc.save(filename)

def upload_file():
    # Function to handle file upload
    global file_path
    file_path = read_excel_file()

def logo_upload():
    # Function to handle file upload
    global logo_path
    logo_path = read_excel_file()

def create_invitations():
    # Function to handle the 'Create' button click event
    session = session_var.get()
    period = period_var.get()
    au = au_entry.get()
    if not file_path:
        messagebox.showerror("Error", "Please upload an Excel file.")
        return
    process_excel(file_path, session, period, au)
    messagebox.showinfo("Success", "Convocations created successfully!")
    
    # Prompt user to save the file
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], title=f"invitations_profs_{today.strftime('%Y_%m_%d')}")
    if save_path:
        save_path = save_path.strip()
        if save_path:
            try:
                import shutil
                # Get the full path of the source file
                source_file = f"invitations_profs_{today.strftime('%Y_%m_%d')}.docx"
                # Move the file to the specified save path
                shutil.move(source_file, save_path)
                messagebox.showinfo("Success", f"File saved successfully at {save_path}.")
            except Exception as e:
                messagebox.showerror("Error", f"Error saving file: {e}")
# Create main window
root = tk.Tk()
root.title("Covocation Creation")
root.geometry("600x400")  # Set width x height

# Heading
heading_label = tk.Label(root, text="Convocation Creation", font=("Arial", 20))
heading_label.pack(pady=10)

# File Upload
file_path = None
file_upload_button = tk.Button(root, text="Upload Excel File", command=upload_file, background="#abcdef")
file_upload_button.pack(pady=5)

# logo Upload
logo_path = None
logo_upload_button = tk.Button(root, text="Upload a FSA Logo", command=logo_upload,background="#fecdab")
logo_upload_button.pack(pady=5)

# Dropdown for Session
session_var = tk.StringVar(root)
session_label = tk.Label(root, text="Session:")
# session_label.grid(column=0, row=0, sticky=tk.W, padx=5, pady=5)
session_var.set("Normale")  # Set default value

session_label.pack()
session_dropdown = tk.OptionMenu(root, session_var, "Rattrappage", "Normale")
session_dropdown.pack()

# Dropdown for Period
period_var = tk.StringVar(root)
period_label = tk.Label(root, text="Period:")
period_var.set("Automne")  # Set default value
period_label.pack()
period_dropdown = tk.OptionMenu(root, period_var, "Automne", "Printemps")
period_dropdown.pack()

# Text Entry for AU
au_label = tk.Label(root, text="Année Universitaire (20XX/20YY):")
au_label.pack()
au_entry = tk.Entry(root)
au_entry.pack()

# Submit Button
submit_button = tk.Button(root, text="Generate Convocations", command=create_invitations, background='#abcdef')
submit_button.pack(pady=5)

root.mainloop()
