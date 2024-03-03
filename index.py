import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
# from datetime import date
import datetime

def read_excel_file(file_path):
    """
    Read an Excel file and return the DataFrame.
    """
    return pd.read_excel(file_path, skiprows=3, header=None, index_col=[0], engine='openpyxl')


##### PRETREATMENT #####
def pretreatment(raw_data_df):
    """
    Perform pretreatment on the DataFrame.
    """
    raw_data_df.iloc[7:, 0] = raw_data_df.iloc[7:, 0].ffill()
    raw_data_df.iloc[0, :] = raw_data_df.iloc[0, :].ffill()
    raw_data_df.iloc[1, :] = raw_data_df.iloc[1, :].ffill()
    raw_data_df.to_excel("raw_data_df.xlsx", index=False)
    return raw_data_df.iloc[6:,:]

# processed_data_df.to_excel("processed_data_df.xlsx", index=False)

"""
    @raw_data_df: big dataframe containing original excel file with header=None and 3 rows skipped
    @processed_data_df: Region of interest extracted from raw_data_df 
    @return: a dict grouping each professor along with their schedules(subject, time, date...)
"""
def grouping_profs_info_in_a_dict(raw_data_df:pd.DataFrame, processed_data_df:pd.DataFrame):

    schedule_data = {}
    # print("professor   subject   date   time niveau")
    for x, row in enumerate(processed_data_df.values):  
        for y, value in enumerate(row):    
            if value == "*":
                professor = processed_data_df.iloc[x, 0]
                date = raw_data_df.iloc[0, y]
                time = raw_data_df.iloc[1, y]
                niveau = raw_data_df.iloc[3, y]
                subject = raw_data_df.iloc[4, y]
                # print(f"{professor}\t{subject}\t{date}\t{time}\t{niveau}")
                schedule_info = {'subject': subject, 'date': date, 'time': time,'niveau':niveau}
                if professor in schedule_data:
                    schedule_data[professor].append(schedule_info)
                else:
                    schedule_data[professor] = [schedule_info]
    return schedule_data

"""
    @schedule_data: dict containing all info about profs schedules extracted from dataframe
    @return: generate a docx file 
"""
def generate_docx(schedule_data:dict, session="Normale", periode="Printemps",au="2023/2024"):
    """
    Generate a DOCX document from schedule data.
    """
    doc = Document()
    today = datetime.date.today()
    for professor, info_list in schedule_data.items():
        # Add professor's name as heading
        # doc.add_picture('endark.png', width=Inches(4))
        logo_paragraph = doc.add_paragraph()

    # Add the logo to the paragraph
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture('endark.png', width=Inches(4))  # Adjust the width as needed

        # Set the alignment of the parag¨raph to centered
        logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()
        doc.add_paragraph()

        doc.add_paragraph(f"A Mme/ Mr: {professor}")
        # doc.add_heading(f"Professor: {professor}", level=1)
        doc.add_paragraph()

        doc.add_paragraph(f"Objet: Convocation aux surveillances des Examens")
        doc.add_paragraph("Cher(e) collègue,")
        # Create a table for the subjects
        doc.add_paragraph(f" Nous vous saurions gé de bien vouloir prendre toutes les dispositions nécessaires pour assurer la surveillance des épreuves écrites de la session {session} de {periode} {au} aux jours et horaires indiqués ci-dessus:")
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Subject'
        hdr_cells[1].text = 'Date'
        hdr_cells[2].text = 'Time'
        hdr_cells[3].text = 'Niveau'
        # hdr_cells[4].text = 'Local'
        
        # Add each subject, date, and time as a row in the table
        for info in info_list:
            subject = info['subject']
            date = info['date']
            time = info['time']
            niveau = info['niveau']
            row_cells = table.add_row().cells
            row_cells[0].text = subject
            row_cells[1].text = date
            row_cells[2].text = time
            row_cells[3].text = niveau
        doc.add_paragraph()

        doc.add_paragraph("Nous vous remercions de votre précieuse collaboration")
        # Add a page break between professors
        doc.add_paragraph(f"Ait Melloul le: {today.strftime('%d-%m-%Y')} ").alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
        doc.add_paragraph()
        doc.add_paragraph(f"Le doyen").alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT

        doc.add_page_break()

    # Save the document
    # today = str(date.today()).replace('-', '_')
    # filename = f"invitations_profs_{today}.docx"
      # Get the current date
    filename = f"invitations_profs_{today.strftime('%Y_%m_%d')}.docx"  # Format date in YYYY-MM-DD format
    doc.save(filename)

def main():

    excel_path = "schedule.xlsx"
    raw_data_df = read_excel_file(file_path=excel_path)
    processed_data_df = pretreatment(raw_data_df=raw_data_df)

    schedule_data = grouping_profs_info_in_a_dict(raw_data_df=raw_data_df, processed_data_df=processed_data_df)
    generate_docx(schedule_data)
if __name__ == "__main__":
    main()